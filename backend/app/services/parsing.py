"""Extract plain text from an uploaded resume (PDF via PyMuPDF, DOCX via python-docx)."""

from __future__ import annotations

import io
import re


class UnsupportedFileError(ValueError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _clean(_pdf_text(data))
    if name.endswith(".docx"):
        return _clean(_docx_text(data))
    raise UnsupportedFileError("Only PDF and DOCX files are supported.")


def _pdf_text(data: bytes) -> str:
    import fitz  # PyMuPDF — imported lazily to keep import cost off the hot path

    parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" · ".join(cells))
    return "\n".join(parts)


def _clean(text: str) -> str:
    """Normalize whitespace: trim lines, collapse runs of blank lines, drop noise."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Normalize bullet glyphs to a plain marker so downstream parsing is uniform.
    text = re.sub(r"[•▪◦‣·]", "-", text)
    lines = [ln.strip() for ln in text.split("\n")]
    out: list[str] = []
    blank = 0
    for ln in lines:
        if not ln:
            blank += 1
            if blank <= 1:
                out.append("")
            continue
        blank = 0
        out.append(ln)
    return "\n".join(out).strip()
